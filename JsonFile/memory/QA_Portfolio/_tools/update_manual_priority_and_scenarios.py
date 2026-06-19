from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
WORKBOOK_PATH = ROOT / "TC_List.xlsx"
TODAY = "2026-06-14"
LIGHT_BLUE = "D9EAF7"
LIGHTER_BLUE = "EAF6FF"
BLACK = "000000"
HEADER_DARK = "1F4D78"
HEADING_BLUE = "2E74B5"

PRIORITY_BY_TC = {
    "TC_001": ("P1", "로비 시작->GameScene 진입은 기본 실행 가능 여부를 판단하는 핵심 Smoke."),
    "TC_002": ("P1", "스토리 진행 불가 시 주요 플레이 루프가 차단됨."),
    "TC_003": ("P4", "패치노트 여백/가독성 중심의 UI 개선 항목."),
    "TC_004": ("P4", "패치노트 갱신 확인은 콘텐츠 표시 품질 항목."),
    "TC_005": ("P3", "패치노트 노출 설정 저장은 UX/설정 저장 품질 항목."),
    "TC_006": ("P2", "아이템 정보 표시 오류는 장비 판단과 UX에 직접 영향."),
    "TC_007": ("P1", "장비 장착 실패는 핵심 성장/전투 준비 흐름을 막음."),
    "TC_008": ("P1", "장비 해제 실패는 인벤토리/상태 복구 결함으로 이어짐."),
    "TC_009": ("P1", "장비 능력치 반영 결함이 이미 관찰되어 회귀 우선순위 높음."),
    "TC_010": ("P1", "장비 부가 효과는 전투 결과와 데이터 연결을 동시에 검증하는 핵심 항목."),
    "TC_011": ("P1", "장비 해제 후 효과 잔존은 전투 밸런스와 회귀 위험이 큼."),
    "TC_012": ("P2", "버프/디버프 UI는 전투 상태 인지에 중요하지만 즉시 진행 차단은 아님."),
    "TC_013": ("P2", "광폭화 장비 연결은 특정 옵션 통합 검증 항목."),
    "TC_014": ("P2", "광폭화 전투 시작 적용은 특정 옵션의 실제 전투 동작 검증."),
    "TC_015": ("P3", "공격 속도 50% 증가는 밸런스 체감/수치 조정 중심."),
    "TC_016": ("P2", "반동 피해는 전투 생존과 옵션 리스크에 직접 영향."),
    "TC_017": ("P2", "방어력 감소/피해 증가 계산은 전투 결과에 직접 영향."),
    "TC_018": ("P1", "광폭화 중첩은 발생 시 스탯 폭증/전투 붕괴 위험이 큼."),
    "TC_019": ("P1", "전투 종료 후 임시 효과 잔존은 다음 전투까지 오염될 수 있음."),
    "TC_020": ("P3", "복수 피해 증가 합산은 확장/엣지 검증이며 추가 데이터 이후 확인."),
    "TC_021": ("P3", "Allocator 로그는 품질 신뢰도 이슈지만 게임 기능 직접 결함은 아님."),
    "TC_022": ("P1", "몬스터 옵션 브릿지 수집은 몬스터 패시브 전체 런타임 연결의 기준."),
    "TC_023": ("P2", "몬스터 출혈/독 OnHit는 전투 상태 이상 체감과 연결됨."),
    "TC_024": ("P2", "강제 빗나감은 전투 판정 체감에 큰 영향을 주는 상태 이상."),
    "TC_025": ("P2", "몬스터 BattleStart 패시브 정리는 전투 단위 상태 오염 방지 항목."),
    "TC_026": ("P3", "복합 효과는 상위 몬스터/보스 확장 검증 항목."),
    "TC_027": ("P1", "몬스터 옵션 데이터 무결성은 잘못된 ID를 전투 전 차단하는 핵심 회귀."),
}


def header_map(ws, row=1) -> dict[str, int]:
    return {str(cell.value).strip(): cell.column for cell in ws[row] if cell.value is not None}


def add_note_suffix(value: str, suffix: str) -> str:
    if suffix in value:
        return value
    return f"{value}\n{suffix}".strip()


def style_range(ws, header_row: int = 1) -> None:
    thin = Side(style="thin", color=BLACK)
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    max_row = ws.max_row
    max_col = ws.max_column
    for row in ws.iter_rows(min_row=header_row, max_row=max_row, max_col=max_col):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = Font(name="맑은 고딕", size=10, color=BLACK)
    for cell in ws[header_row]:
        cell.fill = PatternFill("solid", fgColor=LIGHT_BLUE)
        cell.font = Font(name="맑은 고딕", bold=True, size=10, color=BLACK)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = f"A{header_row + 1}"
    ws.sheet_view.showGridLines = False
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(max_col)}{max_row}"
    for col_idx in range(1, max_col + 1):
        values = [str(ws.cell(row=row, column=col_idx).value or "") for row in range(header_row, min(max_row, 80) + 1)]
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max(len(v) for v in values) + 2, 10), 46)


def add_priority_formatting(ws, priority_col: int, start_row: int, end_row: int) -> None:
    col = get_column_letter(priority_col)
    rng = f"{col}{start_row}:{col}{end_row}"
    ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{start_row},2)="P1"'], font=Font(color="C00000", bold=True)))
    ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{start_row},2)="P2"'], font=Font(color="C55A11", bold=True)))
    ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{start_row},2)="P3"'], font=Font(color="1F4E79")))
    ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{start_row},2)="P4"'], font=Font(color="7F7F7F")))


def update_manual_priorities(wb) -> None:
    for sheet_name in ["수동 플레이 테스트 TC", "수동_TC"]:
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        headers = header_map(ws)
        if "TC ID" not in headers or "우선도" not in headers:
            continue
        note_col = headers.get("이슈/비고")
        priority_col = headers["우선도"]
        for row in range(2, ws.max_row + 1):
            tc_id = ws.cell(row=row, column=headers["TC ID"]).value
            if tc_id not in PRIORITY_BY_TC:
                continue
            priority, reason = PRIORITY_BY_TC[tc_id]
            ws.cell(row=row, column=priority_col).value = priority
            if note_col:
                current = str(ws.cell(row=row, column=note_col).value or "")
                ws.cell(row=row, column=note_col).value = add_note_suffix(current, f"우선도 재분류({TODAY}): {reason}")
        style_range(ws, 1)
        add_priority_formatting(ws, priority_col, 2, ws.max_row)


def recreate_scenario_sheet(wb) -> None:
    name = "시나리오_초안"
    if name in wb.sheetnames:
        del wb[name]
    ws = wb.create_sheet(name)
    rows = [
        [
            "Scenario ID",
            "시나리오명",
            "목적",
            "범위",
            "사전 조건",
            "사용자 흐름",
            "테스트 데이터",
            "확인 포인트",
            "자동화 후보",
            "연결 TC",
            "Evidence",
            "상태",
        ],
        [
            "SCN-001",
            "로비에서 게임 시작 후 첫 스토리 진행",
            "핵심 진입 플로우가 실제 사용자 흐름으로 동작하는지 확인",
            "LobbyScenes -> GameScene -> 첫 스토리 클릭 진행",
            "SaveManager 시작 버튼 연결, Main_Script_Master_Main 로드 가능",
            "로비 진입 / 시작 버튼 클릭 / GameScene 로드 / 첫 스토리 출력 / 화면 클릭으로 다음 스토리 진행",
            "LobbyScenes, GameScene, Main_Script_Master_Main",
            "씬 전환, 텍스트 출력, 콘솔 오류, 페이드 동작",
            "PlayMode Smoke 후보",
            "TC_001, TC_002",
            "07_Evidence/SCN-001/",
            "초안",
        ],
        [
            "SCN-002",
            "장비 지급 후 장착/해제/스탯 반영",
            "장비 데이터와 인벤토리/스탯/옵션 연결을 사용자 관점에서 확인",
            "아이템 정보 표시 -> 장착 -> 스탯 반영 -> 해제 -> 효과 제거",
            "테스트 장비 지급 가능, InventoryManager 정상 동작",
            "장비 지급 / 인벤토리 열기 / 아이템 정보 확인 / 장착 / 스탯 변화 기록 / 해제 / 스탯 복구 확인",
            "Weapon_Master, Armor_Master, Option_Master",
            "장착 슬롯, 인벤토리 이동, 스탯 변화, 옵션 등록/제거",
            "부분 자동화 후보",
            "TC_006~TC_011",
            "07_Evidence/SCN-002/",
            "초안",
        ],
        [
            "SCN-003",
            "전투 상태 이상과 버프 UI 확인",
            "전투 효과가 수치와 UI 양쪽에서 설명 가능한지 확인",
            "장비 옵션/몬스터 패시브 -> 전투 로그 -> UI 표시",
            "전투 진입 가능, 버프 UI/로그 확인 가능",
            "테스트 장비 장착 / 테스트 몬스터 전투 / 공격 적중 / 버프/디버프 표시 / HP 변화 기록",
            "Option_003, Option_004, Option_008~010, 몬스터 패시브 데이터",
            "HP 변화, 상태 아이콘, 발동 타이밍, 로그 가독성",
            "수동 중심",
            "TC_010, TC_012, TC_023, TC_024",
            "07_Evidence/SCN-003/",
            "초안",
        ],
        [
            "SCN-004",
            "상점 가격/구매 체감 검증",
            "아이템 가격이 획득 재화와 진행 단계에 맞는지 검토",
            "상점 목록 -> 가격 확인 -> 구매 가능 시점 추정",
            "상점 UI 접근 가능, BlackSmith.json 최신 상태",
            "상점 진입 / 아이템 목록 확인 / 가격 기록 / 현재 획득 골드 기준 구매 가능 시점 계산 / 체감 메모 작성",
            "BlackSmith.json, Weapon_Master, Armor_Master, 플레이어 골드 획득량",
            "가격 표시, 구매 조건, 등급별 가격 차이, 체감 난이도",
            "수동/데이터 검증 후보",
            "추가 TC 필요",
            "07_Evidence/SCN-004/",
            "초안",
        ],
    ]
    for row in rows:
        ws.append(row)
    style_range(ws, 1)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")
        borders.append(element)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_border(cells[idx])
    doc.add_paragraph()


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 16, HEADING_BLUE),
        ("Heading 2", 13, HEADING_BLUE),
        ("Heading 3", 12, HEADER_DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_doc_appendix(path: Path, heading: str, paragraphs: list[str], tables: list[tuple[list[str], list[list[str]]]]) -> None:
    doc = Document(path)
    style_doc(doc)
    existing_text = "\n".join(p.text for p in doc.paragraphs)
    if heading in existing_text:
        doc.save(path)
        return
    doc.add_page_break()
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    for headers, rows in tables:
        add_table(doc, headers, rows)
    doc.save(path)


def update_docs() -> None:
    add_doc_appendix(
        ROOT / "01_Test_Strategy.docx",
        "수동 TC 우선도 재분류 기준",
        [
            "수동 TC의 우선도는 P1~P4로 표준화한다. P1은 게임 진행/전투/데이터 연결을 막는 핵심 회귀, P2는 주요 기능 품질, P3는 밸런스/도구/확장 확인, P4는 UI 문구와 가벼운 개선 항목이다.",
            "이미지는 TC 본문에 직접 대량 삽입하지 않고 Evidence 경로로 연결한다. 체감/밸런스 메모는 판단과 수치 기록을 남기는 칸으로 유지한다.",
        ],
        [
            (
                ["우선도", "적용 기준", "대표 TC"],
                [
                    ["P1", "진행 차단, 전투/장비/데이터 연결 핵심 결함", "TC_001, TC_002, TC_007~TC_011, TC_018, TC_019, TC_022, TC_027"],
                    ["P2", "주요 기능 품질과 실제 전투 영향", "TC_006, TC_012~TC_014, TC_016, TC_017, TC_023~TC_025"],
                    ["P3", "밸런스, 도구 안정성, 복합/확장 검증", "TC_005, TC_015, TC_020, TC_021, TC_026"],
                    ["P4", "패치노트 표시/콘텐츠 가독성 중심", "TC_003, TC_004"],
                ],
            )
        ],
    )
    add_doc_appendix(
        ROOT / "03_Test_Execution_Report.docx",
        "시나리오 작성 및 Evidence 운영 위치",
        [
            "시나리오 초안은 TC_List.xlsx의 `시나리오_초안` 시트에 먼저 작성한다. 흐름이 확정되면 수동_TC로 분해하고, 자동화 가능한 부분은 자동_PlayMode_TC 또는 회귀_매핑에 연결한다.",
            "전투 밸런스와 상점 가격은 이미지보다 수치 기록을 우선한다. 이미지는 Evidence 경로에 보관하고 TC의 Evidence 칸 또는 이슈/비고에서 연결한다.",
        ],
        [
            (
                ["기록 위치", "용도"],
                [
                    ["시나리오_초안", "로비->전투->상점처럼 여러 TC를 묶는 사용자 흐름 초안"],
                    ["수동_TC", "실제 실행 가능한 단위 테스트 절차와 결과"],
                    ["자동_PlayMode_TC", "반복 실행 가능한 씬/버튼/런타임 smoke 자동화 후보"],
                    ["07_Evidence", "스크린샷, 로그, 영상, 테스트 결과 파일"],
                ],
            )
        ],
    )
    add_doc_appendix(
        ROOT / "04_Automation_Regression_Report.docx",
        "시나리오 자동화 후보 관리",
        [
            "시나리오는 처음부터 전체 E2E 자동화를 목표로 하지 않는다. 안정적인 사용자 흐름을 먼저 수동으로 검증하고, 반복 가치가 높은 부분만 PlayMode 자동화 후보로 승격한다.",
        ],
        [
            (
                ["시나리오", "자동화 판단"],
                [
                    ["로비->게임씬 진입", "현재 PlayMode Smoke로 유지"],
                    ["장비 장착/해제", "오브젝트와 UI 안정성 확보 후 부분 자동화"],
                    ["전투 밸런스", "수치/로그 중심 수동 검증 유지"],
                    ["상점 가격", "데이터 검증 + 수동 체감 기록 유지"],
                ],
            )
        ],
    )


def main() -> None:
    wb = load_workbook(WORKBOOK_PATH)
    update_manual_priorities(wb)
    recreate_scenario_sheet(wb)
    try:
        wb.save(WORKBOOK_PATH)
        saved_path = WORKBOOK_PATH
    except PermissionError:
        saved_path = WORKBOOK_PATH.with_name("TC_List_priority_scenario_update.xlsx")
        wb.save(saved_path)
    update_docs()
    print(f"updated manual priorities, scenario sheet, and docx appendices: {saved_path}")


if __name__ == "__main__":
    main()
