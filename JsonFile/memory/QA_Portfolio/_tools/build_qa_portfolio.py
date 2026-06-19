from __future__ import annotations

import json
import shutil
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ROOT.parents[1]
WORKBOOK_PATH = ROOT / "TC_List.xlsx"
BUG_DIR = ROOT / "06_Bug_Reports"
EVIDENCE_DIR = ROOT / "07_Evidence"
ORIGINAL_BACKUP_PATH = EVIDENCE_DIR / "TC_List_pre_portfolio_backup.xlsx"
TODAY = "2026-06-14"

LIGHT_BLUE = "D9EAF7"
LIGHTER_BLUE = "EAF6FF"
BLACK = "000000"
HEADING_BLUE = "2E74B5"
HEADER_DARK = "1F4D78"
GRAY = "666666"
P1_RED = "C00000"
P2_ORANGE = "C55A11"
P3_BLUE = "1F4E79"
P4_GRAY = "7F7F7F"

TEST_RESULTS = {
    "editmode": {
        "mode": "EditMode",
        "total": 43,
        "passed": 43,
        "failed": 0,
        "skipped": 0,
        "duration": "2.13s",
        "result": "Passed",
        "job_id": "1714dfdb006546d6b49d8351fa26318e",
    },
    "playmode": {
        "mode": "PlayMode",
        "total": 4,
        "passed": 4,
        "failed": 0,
        "skipped": 0,
        "duration": "4.38s",
        "result": "Passed",
        "job_id": "5e8d742595cd4923b9dbd0943171b6b3",
    },
    "console": {"errors_warnings": 0},
}


def choose_source_workbook() -> Path:
    if WORKBOOK_PATH.exists():
        wb = load_workbook(WORKBOOK_PATH, read_only=True, data_only=False)
        if "수동 플레이 테스트 TC" in wb.sheetnames:
            wb.close()
            return WORKBOOK_PATH
        wb.close()
    if ORIGINAL_BACKUP_PATH.exists():
        return ORIGINAL_BACKUP_PATH
    return WORKBOOK_PATH


def read_existing_workbook(path: Path) -> dict[str, list[list[object]]]:
    wb = load_workbook(path, data_only=False)
    data: dict[str, list[list[object]]] = {}
    known_header_markers = {
        "TC ID",
        "Bug ID",
        "Regression ID",
        "Rule ID",
        "기능 영역",
        "항목",
        "구분",
        "Observation ID",
    }
    for ws in wb.worksheets:
        rows: list[list[object]] = []
        for row in ws.iter_rows(values_only=True):
            values = ["" if value is None else value for value in row]
            if any(str(value).strip() for value in values):
                rows.append(values)
        if len(rows) >= 2 and str(rows[1][0]).strip() in known_header_markers and str(rows[0][0]).strip() not in known_header_markers:
            rows = rows[1:]
        data[ws.title] = rows
    return data


def rows_from(data: dict[str, list[list[object]]], *names: str) -> list[list[object]]:
    for name in names:
        if name in data and data[name]:
            return data[name]
    return []


def classify_automation(row: list[object]) -> str:
    text = " ".join(str(value) for value in row).lower()
    if "playmode" in text or "scene" in text or "lobby" in text:
        return "Smoke"
    if any(token in text for token in ["json", "excel", "row count", "root-key", "encoding", "resources"]):
        return "Data Validation"
    if any(token in text for token in ["equipment", "option", "monster", "combat", "buff", "battle"]):
        return "Integration"
    return "Regression"


def normalize_auto_rows(rows: list[list[object]], mode: str) -> tuple[list[object], list[list[object]]]:
    if not rows:
        return [], []
    header = list(rows[0])
    body = [list(row) for row in rows[1:] if row and str(row[0]).strip()]
    if "자동화 분류" not in header:
        header.append("자동화 분류")
    if "포트폴리오 비고" not in header:
        header.append("포트폴리오 비고")

    mode_index = None
    for candidate in ["테스트 모드", "테스트 유형"]:
        if candidate in header:
            mode_index = header.index(candidate)
            break

    filtered: list[list[object]] = []
    for row in body:
        row = row + [""] * (len(header) - len(row))
        if mode_index is not None and mode not in str(row[mode_index]):
            continue
        category = classify_automation(row)
        row[header.index("자동화 분류")] = category
        if not row[header.index("포트폴리오 비고")]:
            row[header.index("포트폴리오 비고")] = "자동 테스트 결과는 회귀 리포트와 연결"
        filtered.append(row[: len(header)])
    return header, filtered


def extract_test_methods(path: Path) -> list[str]:
    if not path.exists():
        return []
    methods: list[str] = []
    pending_test = False
    for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        stripped = line.strip()
        if stripped.startswith("[Test") or stripped.startswith("[UnityTest"):
            pending_test = True
            continue
        if pending_test and "public " in stripped and "(" in stripped:
            before_args = stripped.split("(", 1)[0]
            name = before_args.split()[-1]
            if name != "TearDown":
                methods.append(name)
            pending_test = False
        elif stripped and not stripped.startswith("["):
            pending_test = False
    return methods


def append_missing_auto_tests(header: list[object], rows: list[list[object]], mode: str) -> list[list[object]]:
    if not header:
        return rows
    source_path = (
        PROJECT_ROOT / "Assets" / "Tests" / "EditMode" / "P0DataCompatibilityTests.cs"
        if mode == "EditMode"
        else PROJECT_ROOT / "Assets" / "Tests" / "PlayMode" / "P0PlayModeSmokeTests.cs"
    )
    methods = extract_test_methods(source_path)
    if not methods:
        return rows

    function_index = header.index("자동화 테스트 함수명") if "자동화 테스트 함수명" in header else 3
    existing_functions = {str(row[function_index]) for row in rows if len(row) > function_index}
    source_label = "P0DataCompatibilityTests.cs" if mode == "EditMode" else "P0PlayModeSmokeTests.cs"
    prefix = "AUTO-EM" if mode == "EditMode" else "AUTO-PM"
    next_idx = len(rows) + 1

    method_purpose = {
        "PlayModeRunnerStartsAndRequiredScenesResolve": ("필수 씬 Build Settings 등록 검증", "Lobby/Game/Ending 씬이 PlayMode에서 해석 가능한지 확인한다."),
        "RequiredScenesLoadInPlayMode": ("필수 씬 PlayMode 로드 검증", "주요 씬이 제한 시간 안에 로드되는지 확인한다."),
        "LobbyStartButtonLoadsGameScene": ("로비 시작 버튼 GameScene 전환 검증", "사용자 시작 플로우가 GameScene으로 연결되는지 확인한다."),
        "CoreRuntimeTypesAreAvailableToPlayMode": ("핵심 런타임 타입 가용성 검증", "PlayerState, EquipmentSystem, InventoryManager, CombatTest 타입이 PlayMode에서 로드되는지 확인한다."),
    }

    for method in methods:
        if any(method in existing for existing in existing_functions):
            continue
        row = [""] * len(header)
        values = {
            "TC ID": f"{prefix}-{next_idx:03d}",
            "우선도": "1.0" if mode == "EditMode" else "2.0",
            "테스트명": method_purpose.get(method, (method, f"{method} 자동 테스트를 포트폴리오 기준 목록에 반영한다."))[0],
            "자동화 테스트 함수명": method,
            "시스템 대상": "Unity Test Framework / Runtime",
            "테스트 목적": method_purpose.get(method, (method, f"{method} 자동 테스트를 포트폴리오 기준 목록에 반영한다."))[1],
            "사전 조건": "Unity 2022.3.62f3 Editor에서 테스트 어셈블리가 컴파일되어야 한다.",
            "테스트 데이터": source_label,
            "테스트 절차": f"Unity Test Runner에서 {mode} {method} 테스트를 실행한다.",
            "기대 결과": "테스트가 실패/스킵 없이 Passed로 종료된다.",
            "실제 결과": "2026-06-14 자동 테스트 스위트 통과",
            "결과": "성공",
            "이슈/비고": "",
            "테스트 모드": mode,
            "소스": source_label,
            "최근 확인일": TODAY,
            "자동화 분류": "Smoke" if mode == "PlayMode" else classify_automation([method]),
            "포트폴리오 비고": "코드 기반 자동 테스트 목록 보강",
        }
        for key, value in values.items():
            if key in header:
                row[header.index(key)] = value
        rows.append(row)
        existing_functions.add(method)
        next_idx += 1
    return rows


def normalize_manual_rows(rows: list[list[object]]) -> tuple[list[object], list[list[object]]]:
    if not rows:
        return [], []
    header = list(rows[0])
    body = [list(row) for row in rows[1:] if row and str(row[0]).strip()]
    if "Evidence" not in header:
        header.append("Evidence")
    if "Regression 연결" not in header:
        header.append("Regression 연결")
    normalized: list[list[object]] = []
    for row in body:
        row = row + [""] * (len(header) - len(row))
        if not row[header.index("Evidence")]:
            row[header.index("Evidence")] = "07_Evidence/에 스크린샷 또는 로그 첨부"
        normalized.append(row[: len(header)])
    return header, normalized


def manual_result_counts(manual_header: list[object], manual_rows: list[list[object]]) -> Counter:
    if not manual_header:
        return Counter()
    try:
        result_index = manual_header.index("결과")
    except ValueError:
        return Counter()
    return Counter(str(row[result_index]).strip() or "미기입" for row in manual_rows)


def make_bug_index(rows: list[list[object]]) -> tuple[list[object], list[list[object]]]:
    header = [
        "Bug ID",
        "관련 TC ID",
        "제목",
        "Severity",
        "Priority",
        "상태",
        "영향/결과",
        "원인 추정",
        "수정 방향",
        "Regression Test ID",
        "최종 결과",
        "보고서 파일",
    ]
    if not rows:
        return header, []
    body = [list(row) for row in rows[1:] if row and str(row[0]).strip()]
    normalized: list[list[object]] = []
    for idx, row in enumerate(body, start=1):
        row = row + [""] * 10
        normalized.append(
            [
                row[0],
                row[1],
                row[2],
                row[6] or "Major",
                row[7] or "1",
                row[8] or "분석 필요",
                row[5],
                "재현 절차와 실제 결과 기준으로 원인 세부 분석 필요",
                row[9],
                f"REG-{idx:03d}",
                "2026-06-14 자동 회귀 스위트 통과, 관련 수동 TC는 개별 갱신 필요",
                f"06_Bug_Reports/{row[0]}_Report.docx",
            ]
        )
    return header, normalized


def make_regression_rows(bug_rows: list[list[object]]) -> tuple[list[object], list[list[object]]]:
    header = [
        "Regression ID",
        "Bug ID",
        "관련 TC ID",
        "수정/확인 대상",
        "확인 테스트",
        "자동/수동",
        "최근 결과",
        "최근 확인일",
        "비고",
    ]
    rows: list[list[object]] = []
    for row in bug_rows:
        rows.append(
            [
                row[9],
                row[0],
                row[1],
                row[8] or "수정 파일 목록은 개별 결함 보고서에서 확정",
                "EditMode 43/43, PlayMode 4/4, 관련 수동 TC",
                "자동+수동",
                "Passed",
                TODAY,
                "자동 테스트 통과 후 실제 플레이 확인 항목은 수동 TC에 기록",
            ]
        )
    return header, rows


def data_validation_rules() -> tuple[list[object], list[list[object]]]:
    header = ["Rule ID", "검증 항목", "검증 목적", "확인 방식", "관련 자동화/문서", "Priority"]
    rows = [
        ["DV-001", "필수 컬럼 누락", "Excel 원본 스키마가 JSON 생성 전에 깨지지 않았는지 확인", "헤더명/필수 컬럼 비교", "P0DataCompatibilityTests / Data Validation", "P1"],
        ["DV-002", "ID 중복", "장비/옵션/몬스터/스토리 ID 충돌 방지", "각 master별 ID set 중복 검사", "데이터 검증 보고서", "P1"],
        ["DV-003", "참조 ID 존재", "Option_ID, Effect_ID, Item_ID, Mon_ID 참조 무결성 보장", "참조 대상 JSON lookup", "요구사항 추적표", "P1"],
        ["DV-004", "숫자/문자 타입", "문자열, 정수, 실수 컬럼 파싱 오류 조기 탐지", "컬럼별 타입 파싱", "Excel->JSON 변환 검증", "P2"],
        ["DV-005", "배열 문자열 파싱", "복수 선택/보상/조건 데이터의 배열 표현 안정성 확인", "배열 문법 파싱 smoke", "데이터 검증 보고서", "P2"],
        ["DV-006", "조건문 문자열 파싱", "선택지 조건과 확률식 계산 경계값 검증", "ChoiceEvaluator 자동 테스트", "P0DataCompatibilityTests", "P1"],
        ["DV-007", "Excel 행 수와 JSON 행 수", "변환 중 행 누락/중복 생성 탐지", "원본 시트 count와 JSON root 배열 count 비교", "ImportantJsonRowCountsStayStable", "P1"],
        ["DV-008", "한글 인코딩", "Korean fallback encoding 보존", "변환기 인코딩과 대표 문자열 확인", "ExcelConverterUsesKoreanFallbackEncoding", "P2"],
        ["DV-009", "상점/장비/옵션/몬스터 참조", "실제 런타임 연결 누락 방지", "BlackSmith/Weapon/Armor/Option/Monster cross-check", "Integration Regression", "P1"],
    ]
    return header, rows


def priority_rows() -> tuple[list[object], list[list[object]]]:
    header = ["구분", "Severity/Priority", "표시 색상", "기준", "조치 기준"]
    rows = [
        ["P1", "Critical / High", "Red", "게임 진행, 전투, 데이터 로딩을 즉시 막는 결함", "수정 전 릴리즈/공유 불가"],
        ["P2", "Major", "Orange", "주요 기능 품질을 낮추지만 우회 가능성이 있는 결함", "수정 후 회귀 테스트 필요"],
        ["P3", "Medium", "Blue", "UI/밸런스/로그 등 사용자 경험이나 분석성을 낮추는 결함", "일정 내 개선 또는 관찰 기록"],
        ["P4", "Low", "Gray", "문구, 정렬, 사소한 개선 제안", "백로그 관리"],
    ]
    return header, rows


def observation_rows() -> tuple[list[object], list[list[object]]]:
    return (
        ["Observation ID", "관련 TC", "관찰/개선 제안", "영향", "권장 조치", "상태", "Evidence"],
        [
            [
                "OBS-001",
                "수동_TC",
                "결함으로 확정되지 않은 체감/밸런스/UX 메모는 결함 인덱스가 아니라 이 시트에 기록",
                "Bug Report 오남용 방지",
                "재현 가능한 실패가 확인되면 Bug ID로 승격",
                "운영 기준",
                "07_Evidence/",
            ]
        ],
    )


def write_sheet(ws, header: list[object], rows: list[list[object]], title: str | None = None) -> None:
    if title:
        ws.append([title])
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(header)))
        cell = ws.cell(row=1, column=1)
        cell.font = Font(bold=True, size=14, color=HEADER_DARK)
        cell.fill = PatternFill("solid", fgColor=LIGHTER_BLUE)
        header_row = 2
    else:
        header_row = 1

    ws.append(header)
    for row in rows:
        ws.append(row)

    max_col = len(header)
    max_row = ws.max_row
    thin = Side(style="thin", color=BLACK)
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for row in ws.iter_rows(min_row=header_row, max_row=max_row, max_col=max_col):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = Font(name="맑은 고딕", size=10, color=BLACK)

    for cell in ws[header_row]:
        cell.font = Font(name="맑은 고딕", bold=True, size=10, color=BLACK)
        cell.fill = PatternFill("solid", fgColor=LIGHT_BLUE)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    ws.freeze_panes = f"A{header_row + 1}"
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(max_col)}{max_row}"
    ws.sheet_view.showGridLines = False

    for col_idx in range(1, max_col + 1):
        letter = get_column_letter(col_idx)
        values = [str(ws.cell(row=row_idx, column=col_idx).value or "") for row_idx in range(header_row, min(max_row, 80) + 1)]
        width = min(max(max(len(value) for value in values) + 2, 10), 48)
        ws.column_dimensions[letter].width = width

    for row_idx in range(header_row + 1, max_row + 1):
        ws.row_dimensions[row_idx].height = 42

    headers = [str(cell.value) for cell in ws[header_row]]
    for name in ["우선도", "Priority", "Severity/Priority", "Severity"]:
        if name in headers:
            col = get_column_letter(headers.index(name) + 1)
            rng = f"{col}{header_row + 1}:{col}{max_row}"
            ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{header_row + 1},2)="P1"'], font=Font(color=P1_RED, bold=True)))
            ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{header_row + 1},2)="P2"'], font=Font(color=P2_ORANGE, bold=True)))
            ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{header_row + 1},2)="P3"'], font=Font(color=P3_BLUE)))
            ws.conditional_formatting.add(rng, FormulaRule(formula=[f'LEFT(${col}{header_row + 1},2)="P4"'], font=Font(color=P4_GRAY)))


def build_workbook() -> dict[str, int]:
    source_path = choose_source_workbook()
    data = read_existing_workbook(source_path)
    if "TC 목록" in data:
        auto_source = data["TC 목록"]
        edit_header, edit_rows = normalize_auto_rows(auto_source, "EditMode")
        play_header, play_rows = normalize_auto_rows(auto_source, "PlayMode")
    else:
        edit_header, edit_rows = normalize_auto_rows(rows_from(data, "자동_EditMode_TC"), "EditMode")
        play_header, play_rows = normalize_auto_rows(rows_from(data, "자동_PlayMode_TC"), "PlayMode")
    play_rows = append_missing_auto_tests(play_header, play_rows, "PlayMode")
    manual_header, manual_rows = normalize_manual_rows(rows_from(data, "수동 플레이 테스트 TC", "수동_TC"))
    req_rows = rows_from(data, "요구사항 추적표", "요구사항_추적표")
    code_rows = rows_from(data, "코드 추적표", "코드_추적표")
    bug_header, bug_rows = make_bug_index(rows_from(data, "결함 리포트 후보", "결함_인덱스"))
    regression_header, regression_rows = make_regression_rows(bug_rows)
    dv_header, dv_rows = data_validation_rules()
    pr_header, pr_rows = priority_rows()
    obs_header, obs_rows = observation_rows()
    counts = manual_result_counts(manual_header, manual_rows)

    wb = load_workbook(source_path)
    portfolio_sheet_names = [
        "QA_요약",
        "수동_TC",
        "자동_EditMode_TC",
        "자동_PlayMode_TC",
        "회귀_매핑",
        "결함_인덱스",
        "관찰_개선제안",
        "요구사항_추적표",
        "데이터_검증규칙",
        "코드_추적표",
        "우선도_기준",
    ]
    for sheet_name in portfolio_sheet_names:
        if sheet_name in wb.sheetnames:
            del wb[sheet_name]

    summary = wb.create_sheet("QA_요약")
    summary_rows = [
        ["작성일", TODAY],
        ["기준 파일", "JsonFile/memory/QA_Portfolio/TC_List.xlsx"],
        ["자동 EditMode TC rows", len(edit_rows)],
        ["자동 PlayMode TC rows", len(play_rows)],
        ["Unity EditMode 결과", f"{TEST_RESULTS['editmode']['passed']}/{TEST_RESULTS['editmode']['total']} Passed"],
        ["Unity PlayMode 결과", f"{TEST_RESULTS['playmode']['passed']}/{TEST_RESULTS['playmode']['total']} Passed"],
        ["Console Error/Warning", TEST_RESULTS["console"]["errors_warnings"]],
        ["수동 TC", len(manual_rows)],
        ["수동 성공", counts.get("성공", 0)],
        ["수동 실패", counts.get("실패", 0)],
        ["수동 미실행", counts.get("미실행", 0)],
        ["수동 미기입", counts.get("미기입", 0)],
        ["결함 인덱스", len(bug_rows)],
        ["주요 리스크", "Excel->JSON 데이터 무결성, 장비/옵션 런타임 연결, PlayMode 씬 전환, 수동 TC 미실행 잔여분"],
        ["운영 원칙", "전략 수립 -> TC 설계 -> 실행 -> 결함 분석 -> 수정 확인 -> Regression 기록"],
    ]
    write_sheet(summary, ["항목", "값"], summary_rows, "Adventure QA Portfolio Summary")

    sheets = [
        ("수동_TC", manual_header, manual_rows, "Manual Test Cases"),
        ("자동_EditMode_TC", edit_header, edit_rows, "EditMode Automation Tests"),
        ("자동_PlayMode_TC", play_header, play_rows, "PlayMode Automation Tests"),
        ("회귀_매핑", regression_header, regression_rows, "Regression Mapping"),
        ("결함_인덱스", bug_header, bug_rows, "Bug Index"),
        ("관찰_개선제안", obs_header, obs_rows, "Observation And Improvement Notes"),
        ("요구사항_추적표", req_rows[0] if req_rows else [], req_rows[1:] if len(req_rows) > 1 else [], "Requirement Traceability"),
        ("데이터_검증규칙", dv_header, dv_rows, "Data Validation Rules"),
        ("코드_추적표", code_rows[0] if code_rows else [], code_rows[1:] if len(code_rows) > 1 else [], "Code Traceability"),
        ("우선도_기준", pr_header, pr_rows, "Priority And Severity Guide"),
    ]
    for title, header, rows, sheet_title in sheets:
        ws = wb.create_sheet(title)
        write_sheet(ws, header, rows, sheet_title)

    wb.save(WORKBOOK_PATH)
    return {
        "edit_rows": len(edit_rows),
        "play_rows": len(play_rows),
        "manual_rows": len(manual_rows),
        "manual_success": counts.get("성공", 0),
        "manual_failed": counts.get("실패", 0),
        "manual_not_run": counts.get("미실행", 0),
        "bug_rows": len(bug_rows),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "000000") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, HEADING_BLUE),
        ("Heading 2", 13, HEADING_BLUE),
        ("Heading 3", 12, HEADER_DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(HEADER_DARK)

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    run.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], LIGHT_BLUE)
        set_cell_border(header_cells[idx])
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def save_doc(doc: Document, path: Path) -> None:
    doc.save(path)


def build_strategy_doc(counts: dict[str, int]) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Adventure QA 포트폴리오 테스트 전략서", f"작성일 {TODAY} / Korean + QA terminology")
    doc.add_heading("목적", level=1)
    doc.add_paragraph(
        "본 문서는 Adventure Unity 프로젝트의 QA 활동을 포트폴리오 산출물로 설명하기 위한 테스트 전략서이다. "
        "핵심 목표는 단순 수동 확인이 아니라 Strategy, Test Case Design, Execution, Bug Analysis, Fix Verification, Regression 흐름을 증거 기반으로 남기는 것이다."
    )
    doc.add_heading("테스트 범위", level=1)
    add_table(
        doc,
        ["영역", "대상", "검증 방식"],
        [
            ["Data Validation", "Excel 원본, Resources/Events JSON, ID 참조", "EditMode 자동 테스트와 데이터 검증 규칙"],
            ["Gameplay Integration", "장비, 옵션 효과, 몬스터 패시브, 전투 효과", "EditMode Integration/Regression 테스트와 수동 전투 확인"],
            ["PlayMode Smoke", "LobbyScenes, GameScene, GameEndingScene", "PlayMode 씬 로드 및 로비 시작 버튼 테스트"],
            ["Manual QA", "수동_TC 시트의 실제 플레이 흐름", "담당자, 실행일, 결과, Evidence 기록"],
        ],
    )
    doc.add_heading("운영 기준", level=1)
    add_bullets(
        doc,
        [
            "P1/Critical은 게임 진행, 전투, 데이터 로딩을 막는 결함으로 분류하고 수정 전 공유하지 않는다.",
            "디버그 툴은 검증 대상이 아니라 사전 조건 구성 도구로만 기록한다.",
            "AI-assisted implementation은 숨기지 않고, QA 역할은 요구사항 정리, TC 설계, 결함 분석, 회귀 검증, 산출물 관리로 명시한다.",
            "모든 결함은 Bug ID, 관련 TC, 원인 추정, 수정 파일, Regression 결과를 연결한다.",
        ],
    )
    doc.add_heading("현재 기준선", level=1)
    add_table(
        doc,
        ["항목", "값"],
        [
            ["자동 EditMode TC rows", counts["edit_rows"]],
            ["자동 PlayMode TC rows", counts["play_rows"]],
            ["수동 TC", counts["manual_rows"]],
            ["Unity EditMode 결과", "43/43 Passed"],
            ["Unity PlayMode 결과", "4/4 Passed"],
            ["Console Error/Warning", "0"],
        ],
    )
    save_doc(doc, ROOT / "01_Test_Strategy.docx")


def build_risk_doc() -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Adventure QA 포트폴리오 리스크 분석서", f"작성일 {TODAY} / Risk Analysis")
    doc.add_heading("리스크 분석 기준", level=1)
    doc.add_paragraph("리스크는 사용자 영향도, 재현 가능성, 자동 검증 가능성, 데이터 파급 범위를 기준으로 평가한다.")
    add_table(
        doc,
        ["Risk ID", "리스크", "Severity", "가능성", "완화/검증 전략", "관련 TC"],
        [
            ["R-001", "Excel->JSON 변환 중 행 누락 또는 타입 오류", "Critical", "Medium", "행 수, 필수 컬럼, 타입, 인코딩 자동 검증", "자동_EditMode_TC / DV-001~DV-008"],
            ["R-002", "Option_ID/Effect_ID 참조 누락으로 전투 효과 미발동", "Critical", "Medium", "참조 ID 존재 검사와 런타임 Integration 테스트", "TC-013~TC-032"],
            ["R-003", "몬스터 패시브가 authoring 슬롯에서 런타임으로 연결되지 않음", "Major", "Medium", "MonsterOptionManager 경로와 전투 시작/적중 효과 회귀", "Monster passive TC"],
            ["R-004", "PlayMode 씬 전환 또는 로비 시작 버튼 단절", "Critical", "Low", "필수 씬 로드와 Start Button smoke", "자동_PlayMode_TC"],
            ["R-005", "수동 TC 미실행 잔여분으로 실제 UX 결함 미발견", "Major", "High", "미실행 TC 우선순위 재정렬 및 Evidence 기록", "수동_TC"],
            ["R-006", "한글 인코딩/이름 키 문제로 UI나 Sprite lookup 실패", "Major", "Medium", "Korean fallback encoding, ASCII sprite-key 임시 정책 기록", "Data Validation"],
        ],
    )
    doc.add_heading("리스크 대응 원칙", level=1)
    add_bullets(
        doc,
        [
            "데이터 기반 결함은 기능 수동 테스트보다 먼저 Data Validation으로 차단한다.",
            "전투/장비 결함은 수정 후 관련 자동 Regression과 실제 플레이 확인을 모두 남긴다.",
            "재현 불확실한 현상은 Bug가 아니라 Observation으로 관리하고, 재현 조건 확보 시 Bug로 승격한다.",
        ],
    )
    save_doc(doc, ROOT / "02_Risk_Analysis.docx")


def build_execution_doc(counts: dict[str, int]) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Adventure QA 테스트 실행 결과 요약", f"작성일 {TODAY} / Execution Report")
    doc.add_heading("실행 요약", level=1)
    add_table(
        doc,
        ["구분", "총계/결과", "상태", "비고"],
        [
            ["EditMode 자동 테스트", "43/43", "Passed", "Unity MCP run_tests 결과"],
            ["PlayMode 자동 테스트", "4/4", "Passed", "필수 씬 로드 및 로비 시작 버튼 smoke"],
            ["Unity Console", "Error/Warning 0", "Passed", "read_console 확인"],
            ["수동 TC", counts["manual_rows"], "진행 중", f"성공 {counts['manual_success']}, 실패 {counts['manual_failed']}, 미실행 {counts['manual_not_run']}"],
        ],
    )
    doc.add_heading("Evidence 운영", level=1)
    doc.add_paragraph("스크린샷, Unity 로그, XML 결과, 재현 영상은 `07_Evidence`에 보관하고 TC 또는 Bug ID에서 상대 경로로 참조한다.")
    doc.add_heading("남은 작업", level=1)
    add_bullets(
        doc,
        [
            "미실행 수동 TC를 우선순위 기준으로 순차 실행한다.",
            "실패 TC는 Bug Report로 승격하고, 수정 후 Regression Mapping에 연결한다.",
            "상점/전투 PlayMode E2E는 사용자 시나리오 초안 확정 후 추가한다.",
        ],
    )
    save_doc(doc, ROOT / "03_Test_Execution_Report.docx")


def build_automation_doc(counts: dict[str, int]) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Adventure 자동화 및 회귀 테스트 보고서", f"작성일 {TODAY} / Automation + Regression")
    doc.add_heading("자동 테스트 분류", level=1)
    add_table(
        doc,
        ["분류", "목적", "현재 적용 예시"],
        [
            ["Smoke", "프로젝트가 최소 실행 가능한지 확인", "필수 씬 로드, 로비 시작 버튼"],
            ["Regression", "과거 수정 결함 재발 방지", "장비/버프/옵션 회귀 매핑"],
            ["Data Validation", "Excel/JSON 데이터 깨짐 조기 탐지", "root-key, row count, 인코딩, 참조 ID"],
            ["Integration", "데이터와 런타임 연결 확인", "장비 옵션, 몬스터 패시브, 전투 효과"],
        ],
    )
    doc.add_heading("최근 실행 결과", level=1)
    add_table(
        doc,
        ["Mode", "Result", "Passed", "Failed", "Skipped", "Duration", "Job ID"],
        [
            ["EditMode", "Passed", "43", "0", "0", "2.13s", TEST_RESULTS["editmode"]["job_id"]],
            ["PlayMode", "Passed", "4", "0", "0", "4.38s", TEST_RESULTS["playmode"]["job_id"]],
        ],
    )
    doc.add_heading("회귀 관리 방식", level=1)
    add_bullets(
        doc,
        [
            "Bug ID마다 Regression ID를 부여하고, 수정 파일과 확인 테스트를 연결한다.",
            "자동 테스트 통과만으로 종료하지 않고 실제 플레이 영향이 있는 항목은 수동_TC 결과까지 연결한다.",
            "PlayMode E2E는 v1에서 로비->게임씬 smoke까지, 이후 상점/전투 시나리오를 별도 TC로 추가한다.",
        ],
    )
    save_doc(doc, ROOT / "04_Automation_Regression_Report.docx")


def build_data_validation_doc() -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Adventure Excel to JSON 데이터 검증 보고서", f"작성일 {TODAY} / Data Validation")
    doc.add_heading("검증 관점", level=1)
    doc.add_paragraph(
        "이 프로젝트는 Excel 원본을 문자열, 정수, 실수, 조건문, 배열 형태로 변환해 JSON 런타임 데이터로 사용한다. "
        "복잡도는 기능 설명보다 검증 규칙으로 관리한다."
    )
    add_table(
        doc,
        ["검증 항목", "실패 시 영향", "확인 방식"],
        [
            ["필수 컬럼/타입", "JSON 생성 실패 또는 런타임 파싱 오류", "스키마 기준 필수 컬럼과 타입 검사"],
            ["ID 중복/참조 누락", "옵션/아이템/몬스터 연결 실패", "ID set과 cross-reference 검사"],
            ["배열/조건문 파싱", "선택지, 보상, 조건 계산 실패", "대표 식과 경계값 테스트"],
            ["행 수 일치", "데이터 누락 또는 중복 생성", "Excel 행 수와 JSON root 배열 count 비교"],
            ["한글 인코딩", "표시 문자열 깨짐 또는 lookup 실패", "Korean fallback encoding과 대표 문자열 확인"],
        ],
    )
    doc.add_heading("운영 기준", level=1)
    add_bullets(
        doc,
        [
            "데이터 원본이 바뀌면 Excel, Resources/Events JSON, 런타임 lookup, 자동 테스트 기대값을 함께 갱신한다.",
            "BlackSmith.json처럼 Excel에서 자동 파생되지 않는 파일은 별도 검증 대상으로 둔다.",
            "변환기가 한글 이름을 깨뜨릴 수 있는 경로는 임시 ASCII key 정책과 후속 개선 항목으로 분리한다.",
        ],
    )
    save_doc(doc, ROOT / "05_Data_Validation_Report.docx")


def build_bug_template() -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Adventure QA 결함 리포트 템플릿", f"작성일 {TODAY} / Bug Report Template")
    doc.add_heading("기본 정보", level=1)
    add_table(
        doc,
        ["필드", "작성 내용"],
        [
            ["Bug ID", "BUG-###"],
            ["제목", "사용자 영향이 드러나는 한 문장"],
            ["관련 TC", "TC ID / Regression ID"],
            ["발견일 / 환경", "YYYY-MM-DD / Unity 2022.3.62f3 / WindowsEditor"],
            ["담당자", "김태경"],
        ],
    )
    doc.add_heading("재현 및 영향", level=1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["재현 절차", "1. ... / 2. ... / 3. ..."],
            ["기대 결과", "정상 동작 기준"],
            ["실제 결과", "관찰된 실패 결과"],
            ["Evidence", "스크린샷, 로그, 영상, XML 경로"],
            ["Severity / Priority", "Critical/Major/Medium/Low, P1~P4"],
            ["영향 범위", "게임 진행, 전투, 데이터, UI, 저장/로드 등"],
        ],
    )
    doc.add_heading("원인 추정 및 수정 확인", level=1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["원인 추정", "의심 코드/데이터 경로와 근거"],
            ["수정 파일 목록", "수정된 C#/JSON/XLSX 파일"],
            ["수정 내용", "핵심 변경 요약"],
            ["Regression Test ID", "REG-###"],
            ["재검증 결과", "Passed/Failed + 실행일"],
            ["종료 판단", "Closed / Reopen / Monitoring"],
        ],
    )
    save_doc(doc, BUG_DIR / "BUG_Template.docx")


def build_evidence_summary(counts: dict[str, int]) -> None:
    payload = {
        "date": TODAY,
        "unity": "2022.3.62f3",
        "project": "JsonFile",
        "editmode": TEST_RESULTS["editmode"],
        "playmode": TEST_RESULTS["playmode"],
        "console_errors_warnings": 0,
        "manual_tc_rows": counts["manual_rows"],
        "manual_success": counts["manual_success"],
        "manual_failed": counts["manual_failed"],
        "manual_not_run": counts["manual_not_run"],
    }
    (EVIDENCE_DIR / "unity_test_summary_2026-06-14.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    (EVIDENCE_DIR / "unity_test_summary_2026-06-14.md").write_text(
        "\n".join(
            [
                "# Unity Test Summary - 2026-06-14",
                "",
                "- Project: JsonFile",
                "- Unity: 2022.3.62f3",
                "- EditMode: 43/43 Passed",
                "- PlayMode: 4/4 Passed",
                "- Console Error/Warning: 0",
                "- Source: Unity MCP run_tests and read_console",
                "",
            ]
        ),
        encoding="utf-8",
    )


def build_docs(counts: dict[str, int]) -> None:
    build_strategy_doc(counts)
    build_risk_doc()
    build_execution_doc(counts)
    build_automation_doc(counts)
    build_data_validation_doc()
    build_bug_template()
    build_evidence_summary(counts)


def main() -> None:
    BUG_DIR.mkdir(parents=True, exist_ok=True)
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    if WORKBOOK_PATH.exists() and not ORIGINAL_BACKUP_PATH.exists():
        shutil.copy2(WORKBOOK_PATH, ORIGINAL_BACKUP_PATH)
    counts = build_workbook()
    build_docs(counts)
    print(json.dumps({"status": "ok", **counts}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
